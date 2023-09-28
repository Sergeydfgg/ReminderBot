import pandas as pd
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
import datetime
import docx
import os
import threading
import telebot
from telebot import types

TOKEN = ''
tb = telebot.TeleBot(TOKEN)


def make_duration_plot(data: list, data_len: int) -> None:
    x_ax = np.arange(0.0, data_len)

    fig, ax = plt.subplots()
    ax.plot(x_ax, data)

    ax.set(xlabel='Кол-во задач', ylabel='Длительность (мин)',
           title='Длительность выполнения задач за выбранный период')
    ax.grid()

    fig.savefig("plt_1.png")


def make_compare_plot(cancel_data_len: int, done_data_len: int, waste_data_len: int) -> None:
    fig, ax = plt.subplots()

    names = ['Отмененные', 'Выполненные', 'Просроченные']
    counts = [cancel_data_len, done_data_len, waste_data_len]
    bar_colors = ['tab:red', 'tab:blue', 'tab:orange']

    ax.bar(names, counts, label=names, color=bar_colors)

    ax.set_ylabel('Кол-во задач')
    ax.set_title('Сравнение итогового статуса задач')
    ax.legend(title='Статус задачи')

    fig.savefig("plt_2.png")


def period_calculation(start_date: str, end_date: str) -> list:
    start_date_parts = start_date.split('-')
    end_date_parts = end_date.split('-')
    start_date = datetime.datetime(int(start_date_parts[0]), int(start_date_parts[1]), int(start_date_parts[2]))
    end_date = datetime.datetime(int(end_date_parts[0]), int(end_date_parts[1]), int(end_date_parts[2]))

    date_range = pd.date_range(
        min(start_date, end_date),
        max(start_date, end_date)
    ).strftime('%Y-%m-%d').tolist()

    return date_range


def make_report(final_table_clear: pd.core.frame.DataFrame, deadline_move: float,
                done_tasks: float, waste_tasks: float, cancel_tasks: float) -> None:
    report = docx.Document()

    report.add_heading('Общая информация за выбранный период')
    report.add_paragraph(f'Среднея длительность выполнения задач - '
                         f'{final_table_clear["duration"].mean():0.2f} (минуты)\n'
                         f'Процент перенесенных дедлайнов - {deadline_move:0.2f}%\n'
                         f'Процент выполненных задач - {done_tasks:0.2f}%\n'
                         f'Процент просроченных задач - {waste_tasks:0.2f}%\n'
                         f'Процент отмененных задач - {cancel_tasks:0.2f}%\n')

    report.add_heading('Графики по данным заданного периода')
    report.add_picture('plt_1.png', width=docx.shared.Inches(6.0))
    report.add_picture('plt_2.png', width=docx.shared.Inches(6.0))

    report.save('Report.docx')
    os.remove('plt_1.png')
    os.remove('plt_2.png')


def prepare_data(user_id: int, start_date: str, end_date: str) -> None:
    matplotlib.use('SVG')
    table = pd.read_csv('DataBase/dates.csv', sep=',', encoding='utf-8')
    user_table = table.query(f'user_id == {user_id}')

    date_range = period_calculation(start_date, end_date)

    final_table_with_deadline = user_table.loc[user_table['task_date'].isin(date_range)]
    final_table_clear = final_table_with_deadline.query('status != "Дедлайн перенесен"')

    waste_len = len(final_table_clear.query('status == "Просрочено"'))
    done_len = len(final_table_clear.query('status == "Сделано"'))
    cancel_len = len(final_table_clear.query('status == "Отменено"'))

    make_duration_plot(final_table_clear['duration'].tolist(), len(final_table_clear['duration']))
    make_compare_plot(cancel_len, done_len, waste_len)

    deadline_move = len(final_table_with_deadline.query('status == "Дедлайн перенесен"')) / len(final_table_clear)*100
    waste_tasks = (waste_len / len(final_table_clear)) * 100
    cancel_tasks = (cancel_len / len(final_table_clear)) * 100
    done_tasks = (done_len / len(final_table_clear)) * 100

    make_report(final_table_clear, deadline_move, done_tasks, waste_tasks, cancel_tasks)

    with open('Report.docx', 'rb') as report:
        tb.send_document(user_id, report)


def get_stat(user_id: int, start_date: str, end_date: str) -> None:
    t1 = threading.Thread(target=prepare_data, args=(user_id, start_date, end_date), daemon=True)
    t1.start()
